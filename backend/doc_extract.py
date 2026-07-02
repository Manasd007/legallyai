"""Extract plain text from an uploaded legal document (PDF / DOCX / TXT).

PDF uses PyMuPDF (fitz) with a pypdf fallback — the same approach as the corpus
pipeline. Scanned/image-only PDFs yield little or no text; the caller checks the
length and returns a helpful error rather than feeding the LLM an empty doc.
"""
from __future__ import annotations

import io


class ExtractionError(Exception):
    """Unsupported file type or unreadable document."""


IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff", ".gif")


def _looks_like_image(name: str, data: bytes) -> bool:
    if name.endswith(IMAGE_EXTS):
        return True
    return (
        data[:3] == b"\xff\xd8\xff"  # JPEG
        or data[:8] == b"\x89PNG\r\n\x1a\n"  # PNG
        or data[:4] == b"RIFF"  # WEBP container
        or data[:2] in (b"BM",)  # BMP
    )


def _mime_for(name: str) -> str:
    n = name.lower()
    if n.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if n.endswith(".webp"):
        return "image/webp"
    return "image/png"


def extract_text(filename: str, data: bytes) -> tuple[str, dict]:
    """Return (text, meta). meta = {"ocr": bool} so the caller can tell the user
    when text came from OCR rather than the embedded text layer."""
    name = (filename or "").lower()
    meta = {"ocr": False}

    # Image uploads (photos/scans) -> OCR directly.
    if _looks_like_image(name, data):
        import ocr

        meta["ocr"] = True
        return ocr.ocr_image(data, _mime_for(name)), meta

    if name.endswith(".pdf") or data[:5] == b"%PDF-":
        text = _pdf(data)
        # Scanned/image-only PDFs yield little text -> fall back to OCR.
        from config import get_settings

        if len(text.strip()) < get_settings().ocr_min_chars:
            import ocr

            ocr_text = ocr.ocr_pdf(data)
            if len(ocr_text.strip()) > len(text.strip()):
                meta["ocr"] = True
                return ocr_text, meta
        return text, meta

    if name.endswith(".docx") or data[:2] == b"PK":  # docx is a zip
        try:
            return _docx(data), meta
        except Exception as e:  # noqa: BLE001
            raise ExtractionError(f"Could not read this Word document: {e}")

    if name.endswith((".txt", ".md", ".text")):
        return data.decode("utf-8", errors="ignore").strip(), meta

    # Last resort: try to decode as text.
    try:
        return data.decode("utf-8").strip(), meta
    except Exception:  # noqa: BLE001
        raise ExtractionError(
            "Unsupported file type. Please upload a PDF, image, Word (.docx), or text file."
        )


def _pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF

        with fitz.open(stream=data, filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc).strip()
    except Exception:  # noqa: BLE001 - fall back to pypdf
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            return "\n".join((p.extract_text() or "") for p in reader.pages).strip()
        except Exception as e:  # noqa: BLE001
            raise ExtractionError(f"Could not read this PDF: {e}")


def _docx(data: bytes) -> str:
    from docx import Document  # python-docx

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Include table cells — legal docs often put terms in tables.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
